import gradio as gr
import os
import tempfile

from docx import Document
from reflection_agent import AppState, AGENT

def build_docx_file(text: str, filename: str) -> str:
    """
    Build a .docx file on disk and return its path.
    """
    tmp_dir = tempfile.mkdtemp()
    path = os.path.join(tmp_dir, filename)
    doc = Document()
    for line in text.splitlines():
        doc.add_paragraph(line)
    doc.save(path)
    return path


def build_tex_file(text: str, filename: str) -> str:
    """
    Build a .tex file on disk and return its path.
    """
    tmp_dir = tempfile.mkdtemp()
    path = os.path.join(tmp_dir, filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)
    return path


def parse_cv_file(cv_file):
    """
    Parse CV content from a Gradio File input.
    Supports .docx and .tex.
    """
    if cv_file is None:
        return None, None

    path = getattr(cv_file, "name", cv_file)
    filename = os.path.basename(path)

    if filename.lower().endswith(".docx"):
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        fmt = "docx"
    elif filename.lower().endswith(".tex"):
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        fmt = "tex"
    else:
        return None, None

    meta = {"format": fmt, "filename": filename}
    return meta, text

def run_langgraph_agent(user_input: dict):
    initial_state = AppState.from_user_input(user_input)
    final_state = AGENT.invoke(initial_state)

    cover_letter_text = (
        final_state.get("cover_letter_final")
        or final_state.get("cover_letter_draft", "")
    )
    improved_cv_text = (
        final_state.get("cv_final")
        or final_state.get("cv_draft", "")
    )
    return cover_letter_text, improved_cv_text

def process_submission(job_description, cv_file):
    if not job_description:
        return "Please paste the job description.", None, None

    if cv_file is None:
        return "Please upload your CV in .docx or .tex format.", None, None

    meta, cv_text = parse_cv_file(cv_file)
    if meta is None:
        return "Unsupported CV format. Please upload a .docx or .tex file.", None, None

    user_input = {
        "job_description": job_description,
        "cv_text": cv_text,
        "cv_format": meta["format"],
        "cv_filename": meta["filename"],
    }

    # call compiled LangGraph app
    cover_letter_text, improved_cv_text = run_langgraph_agent(user_input)

    # Build downloadable files
    cover_letter_path = build_docx_file(
        cover_letter_text,
        filename="cover_letter.docx"
    )

    if meta["format"].lower() == "docx":
        refined_cv_path = build_docx_file(
            improved_cv_text,
            filename=meta["filename"],
        )
    else:  # "tex"
        refined_cv_path = build_tex_file(
            improved_cv_text,
            filename=meta["filename"],
        )
    return cover_letter_text, cover_letter_path, refined_cv_path

def main():

    with gr.Blocks() as demo:
        gr.Markdown("# Job Application Assistant")

        with gr.Row():
            # LEFT: INPUTS
            with gr.Column():
                gr.Markdown("## Inputs")

                job_desc = gr.Textbox(
                    label="Job Description",
                    placeholder="Paste the job description here...",
                    lines=14,
                )

                cv_upload = gr.File(
                    label="CV Upload (.docx or .tex)",
                    file_types=[".docx", ".tex"],
                )

                submit_btn = gr.Button("Submit", variant="primary")

            # RIGHT: OUTPUTS
            with gr.Column():
                gr.Markdown("## Outputs")

                cover_letter_box = gr.Textbox(
                    label="Cover Letter (you can edit this)",
                    lines=18,
                    interactive=True,
                )

                download_cl_btn = gr.DownloadButton(
                    label="Download Cover Letter (.docx)",
                    value="cover_letter.docx"
                )

                download_cv_btn = gr.DownloadButton(
                    label="Download Refined CV (original format)",
                    value="refined_cv"
                )

        # Wire submit -> processing + outputs
        submit_btn.click(
            fn=process_submission,
            inputs=[job_desc, cv_upload],
            outputs=[cover_letter_box, download_cl_btn, download_cv_btn],
        )

    demo.launch()

if __name__ == "__main__":
    main()